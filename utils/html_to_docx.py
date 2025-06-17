
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from bs4 import BeautifulSoup, NavigableString, Tag
from io import BytesIO
import base64
import tempfile

def convert_html_to_docx(html: str, config: dict) -> str:
    doc = Document()

    font = config.get("fontFamily", "Arial")
    size = int(config.get("fontSize", "11px").replace("px", ""))
    spacing = float(config.get("lineSpacing", 1.15))
    logo_width = int(config.get("logoSize", "250px").replace("px", ""))

    style = doc.styles['Normal']
    style.font.name = font
    style.font.size = Pt(size)

    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    soup = BeautifulSoup(html, 'html.parser')

    # Logo injection
    logo_tag = soup.find("img", class_="resume-logo")
    if logo_tag and "base64" in logo_tag.get("src", ""):
        base64_data = logo_tag["src"].split("base64,")[1]
        logo_bytes = BytesIO(base64.b64decode(base64_data))
        doc.add_picture(logo_bytes, width=Inches(logo_width / 96))
        doc.add_paragraph()

    last_top_tag = None

    for elem in soup.find_all(["h2", "p", "ul", "li"], recursive=True):
        if isinstance(elem, NavigableString):
            continue

        text = elem.get_text(strip=True)
        if not text:
            continue

        tag = elem.name
        class_list = elem.get("class", [])

        if tag == "h2":
            if last_top_tag and last_top_tag != "h2":
                doc.add_paragraph()  # insert newline before new section
            para = doc.add_paragraph()
            run = para.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(size + 1)
            para.paragraph_format.space_after = Pt(12)
            last_top_tag = "h2"

        elif tag == "p":
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.size = Pt(size)
            para.paragraph_format.line_spacing = spacing

            if "profile-paragraph" in class_list:
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(8)
            elif "profile-paragraph-note" in class_list:
                para.paragraph_format.space_after = Pt(14)
            elif "work-dates" in class_list:
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(2)
            elif "work-title" in class_list:
                para.paragraph_format.space_after = Pt(8)
            elif "edu-line-1" in class_list:
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(4)
            elif "edu-line-2" in class_list:
                para.paragraph_format.space_after = Pt(4)
            elif "info-field" in class_list:
                para.paragraph_format.space_after = Pt(4)
            else:
                para.paragraph_format.space_after = Pt(5)

            last_top_tag = "p"

        elif tag == "li":
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(text)
            run.font.size = Pt(size)
            para.paragraph_format.line_spacing = spacing
            para.paragraph_format.space_after = Pt(5)
            last_top_tag = "li"

    # Final blank line after all content
    doc.add_paragraph()

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return temp_file.name
